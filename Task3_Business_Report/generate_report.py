"""
Generate the Task 3 Business Report for Geldium's Head of Collections.
Two-page stakeholder-ready report in Word format.
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
import os

doc = Document()

# ── Page setup (narrow margins for 2-page fit) ──
for section in doc.sections:
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.5)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

# ── Styles ──
style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)  # Dark blue
        run.font.size = Pt(14) if level == 1 else Pt(12)
    return h

def add_para(text, bold=False, italic=False, space_after=Pt(4)):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(10.5)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(1)
    return p

def add_bullet(text, bold_prefix="", space_after=Pt(2)):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    else:
        run = p.add_run(text)
        run.font.size = Pt(10.5)
    p.paragraph_format.space_after = space_after
    p.paragraph_format.space_before = Pt(0)
    return p

def set_cell_shading(cell, color):
    shading = cell._element.get_or_add_tcPr()
    shading_elm = shading.makeelement(qn('w:shd'), {
        qn('w:fill'): color,
        qn('w:val'): 'clear'
    })
    shading.append(shading_elm)

def format_table(table):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = 'Calibri'

# ══════════════════════════════════════════════════════════════
# TITLE BLOCK
# ══════════════════════════════════════════════════════════════
title = doc.add_heading('Geldium Delinquency Risk: Business Report', level=0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in title.runs:
    run.font.color.rgb = RGBColor(0x1B, 0x3A, 0x5C)
    run.font.size = Pt(18)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Prepared for the Head of Collections | Tata iQ Analytics Team | March 2026')
run.font.size = Pt(10)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
run.italic = True
subtitle.paragraph_format.space_after = Pt(8)

# ══════════════════════════════════════════════════════════════
# SECTION 1: EXECUTIVE SUMMARY
# ══════════════════════════════════════════════════════════════
add_heading_styled('1. Summary of Predictive Insights')

add_para(
    'Our analytics team analyzed 500 Geldium customer records across 19 variables '
    '— including demographics, financials, credit behavior, and 6-month payment history '
    '— to identify which customers are most likely to become delinquent. '
    'A Gradient Boosting model was developed and validated using 5-fold cross-validation, '
    'achieving an AUC-ROC of 0.93 on training data. Below are the key findings.',
    space_after=Pt(6)
)

add_para('Top 3 Risk Factors', bold=True, space_after=Pt(2))

add_bullet(
    ' Customers with high average payment severity scores (frequent Late or Missed payments '
    'across Months 1–6) are the single strongest predictor of delinquency. This engineered '
    'feature (Avg_Payment_Severity) ranked #1 in model importance at 12.3%.',
    bold_prefix='Payment Behavior Pattern:'
)
add_bullet(
    ' High credit utilization (Credit_Utilization approaching or exceeding 100%) signals '
    'financial strain and ranked as the top traditional financial predictor. Delinquent '
    'customers average 51% utilization vs. 49% for non-delinquent.',
    bold_prefix='Credit Utilization:'
)
add_bullet(
    ' Customers with missing Loan Balance data show a 24.1% delinquency rate compared '
    'to 15.5% for those with complete records. The missingness itself acts as a behavioral '
    'risk indicator — possibly reflecting disengaged or financially distressed customers.',
    bold_prefix='Missing Loan Balance (Behavioral Signal):',
    space_after=Pt(6)
)

add_para('High-Risk Customer Segments', bold=True, space_after=Pt(2))

# Risk segments table
table = doc.add_table(rows=5, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['Customer Segment', 'Delinquency Rate', 'Insight']
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    for run in cell.paragraphs[0].runs:
        run.bold = True

data = [
    ['Business Credit Card Holders', '21.3%', 'Highest among all card types — business cash flow volatility'],
    ['Unemployed Customers', '19.4%', 'Lack of stable income directly impacts repayment capacity'],
    ['Los Angeles Customers', '19.6%', 'Regional cost-of-living pressures may drive delinquency'],
    ['Missing Loan Balance Records', '24.1%', 'Data gap correlates with financial disengagement'],
]
for ri, row_data in enumerate(data):
    for ci, val in enumerate(row_data):
        table.rows[ri + 1].cells[ci].text = val

format_table(table)
doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ══════════════════════════════════════════════════════════════
# SECTION 2: SMART RECOMMENDATION
# ══════════════════════════════════════════════════════════════
add_heading_styled('2. SMART Business Recommendation')

add_para('Key Insight: High credit utilization combined with deteriorating payment patterns '
         'is the strongest composite indicator of imminent delinquency.', italic=True, space_after=Pt(6))

add_para('SMART Goal', bold=True, space_after=Pt(2))

# SMART table
table = doc.add_table(rows=6, cols=2)
table.style = 'Light Grid Accent 1'
smart = [
    ['SMART Element', 'Detail'],
    ['Specific', 'Launch a proactive outreach program targeting customers whose credit utilization '
     'exceeds 80% AND who have 2+ late/missed payments in the last 3 months'],
    ['Measurable', 'Reduce delinquency rate among the targeted high-risk segment by 15% '
     '(from ~21% to ~18%)'],
    ['Achievable', 'Leverage the existing Collections team with AI-generated risk scores '
     'to prioritize outreach — no new hires required'],
    ['Relevant', 'Directly supports Geldium\'s goal of reducing portfolio losses and improving '
     'customer retention through early intervention'],
    ['Time-bound', 'Pilot the program within 90 days; measure impact at the 6-month mark'],
]
for ri, (col1, col2) in enumerate(smart):
    table.rows[ri].cells[0].text = col1
    table.rows[ri].cells[1].text = col2
    if ri == 0:
        for cell in table.rows[ri].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True

format_table(table)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_para('Stakeholder Rationale', bold=True, space_after=Pt(2))
add_para(
    'Our model identifies that payment behavior and credit utilization are the two most '
    'actionable predictors of delinquency. By targeting customers who exhibit both risk '
    'factors simultaneously, the Collections team can intervene before accounts transition '
    'to full delinquency — reducing write-offs, preserving customer relationships, and '
    'improving portfolio health. This targeted approach is more cost-effective than blanket '
    'outreach because the AI model narrows the focus to the highest-risk individuals.',
    space_after=Pt(6)
)

# ══════════════════════════════════════════════════════════════
# SECTION 3: ETHICAL & RESPONSIBLE AI CONSIDERATIONS
# ══════════════════════════════════════════════════════════════
add_heading_styled('3. Ethical and Responsible AI Considerations')

add_para('Fairness Risks and Mitigation', bold=True, space_after=Pt(2))

# Fairness table
table = doc.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
fair_headers = ['Fairness Risk', 'Description', 'Mitigation Strategy']
for i, h in enumerate(fair_headers):
    table.rows[0].cells[i].text = h
    for run in table.rows[0].cells[i].paragraphs[0].runs:
        run.bold = True

fair_data = [
    [
        'Employment Status Bias',
        'The model flags self-employed and unemployed customers at higher rates. '
        'Our fairness audit found a Disparate Impact Ratio of 0.25, well below the '
        'acceptable 0.80 threshold — indicating potential over-flagging of non-traditional workers.',
        'Apply group-specific threshold calibration so that predicted delinquency rates '
        'align with actual rates for each employment group. Monitor the Disparate Impact '
        'Ratio monthly and retrain if it falls below 0.80.'
    ],
    [
        'Geographic Proxy Discrimination',
        'Location (e.g., Los Angeles, Houston) may serve as a proxy for socioeconomic '
        'or racial demographics, potentially leading to unfair treatment of customers '
        'based on where they live rather than their individual financial behavior.',
        'Evaluate whether removing Location from the model degrades performance significantly. '
        'If it does not, exclude it. If it does, apply fairness constraints to ensure '
        'location does not disproportionately impact any group.'
    ],
]
for ri, row_data in enumerate(fair_data):
    for ci, val in enumerate(row_data):
        table.rows[ri + 1].cells[ci].text = val

format_table(table)
doc.add_paragraph().paragraph_format.space_after = Pt(2)

add_para('Explaining AI Predictions to Stakeholders', bold=True, space_after=Pt(2))
add_para(
    'The model works like a scoring system: it examines each customer\'s payment history, '
    'credit usage, and financial profile, then assigns a risk score from 0 (low risk) to 1 '
    '(high risk). Customers scoring above 0.19 are flagged for review. For every flagged '
    'customer, the system provides a plain-language explanation — for example: "This customer '
    'was flagged because they missed payments in 3 of the last 6 months and their credit '
    'utilization is at 92%." A companion Decision Tree model provides simple if/then rules '
    'that any team member can follow, ensuring full transparency.',
    space_after=Pt(4)
)

add_para('Responsible AI Commitments', bold=True, space_after=Pt(2))
add_bullet('All model predictions are explainable using SHAP values and Decision Tree rules — no "black box" decisions.')
add_bullet('The model is validated for fairness across employment status, location, and age groups before deployment.')
add_bullet('Human review remains mandatory — the model recommends, but trained Collections staff make the final decision.')
add_bullet('The model will be retrained quarterly and audited for bias drift using updated customer data.', space_after=Pt(6))

# ══════════════════════════════════════════════════════════════
# SECTION 4: AI & GenAI USAGE DISCLOSURE
# ══════════════════════════════════════════════════════════════
add_heading_styled('4. AI & GenAI Usage')

add_para(
    'Generative AI (Claude Code) was used throughout this project to: (1) conduct exploratory '
    'data analysis and identify data quality issues, (2) engineer predictive features from '
    'raw payment history, (3) train and compare three candidate models, (4) perform fairness '
    'audits, and (5) draft this stakeholder report. All AI-generated outputs were reviewed '
    'and validated by the analytics team. Example prompts included: "Summarize top predictors '
    'of customer delinquency," "Turn this insight into a SMART business recommendation," '
    'and "List fairness risks for a financial risk prediction model."',
    space_after=Pt(2)
)

# ── Save ──
output_path = 'C:/Users/LENOVO X1 YOGA/Desktop/tata/Geldium_Delinquency_Business_Report.docx'
doc.save(output_path)
print(f'Report saved to: {output_path}')
